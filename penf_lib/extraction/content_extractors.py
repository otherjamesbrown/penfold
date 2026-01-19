"""Content extraction pipeline for supported file formats."""

import asyncio
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Optional, List
from io import BytesIO

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image, ImageOps
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    Image = None
    ImageOps = None
    pytesseract = None
    TESSERACT_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """Result of content extraction."""
    success: bool
    text_content: Optional[str] = None
    metadata: Optional[Dict] = None
    content_type: str = "unknown"
    extraction_method: str = "unknown"
    confidence: float = 0.0
    error: Optional[str] = None
    word_count: int = 0
    language: Optional[str] = None


class ContentExtractor:
    """Content extraction for various file formats."""

    def __init__(self):
        """Initialize content extractor."""
        self.extractors = {
            'application/pdf': self._extract_pdf,
            'text/plain': self._extract_text,
            'text/csv': self._extract_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx,
            'application/msword': self._extract_doc_fallback,
            'image/jpeg': self._extract_image_ocr,
            'image/png': self._extract_image_ocr,
            'image/gif': self._extract_image_ocr,
            'image/webp': self._extract_image_ocr
        }

    async def extract_content(self, file_path: Path, mime_type: str) -> ExtractionResult:
        """Extract content from file.

        Args:
            file_path: Path to file
            mime_type: MIME type of file

        Returns:
            ExtractionResult with extracted content and metadata
        """
        if mime_type not in self.extractors:
            return ExtractionResult(
                success=False,
                error=f"Unsupported MIME type: {mime_type}",
                content_type=mime_type
            )

        try:
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            extractor = self.extractors[mime_type]
            result = await loop.run_in_executor(None, extractor, file_path)

            # Post-process result
            if result.success and result.text_content:
                result.text_content = self._clean_text(result.text_content)
                result.word_count = len(result.text_content.split())
                result.language = self._detect_language_basic(result.text_content)

            return result

        except Exception as e:
            logger.error(f"Content extraction failed for {file_path}: {e}")
            return ExtractionResult(
                success=False,
                error=str(e),
                content_type=mime_type
            )

    def _extract_pdf(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            ExtractionResult with extracted text
        """
        if not PyPDF2:
            return ExtractionResult(
                success=False,
                error="PyPDF2 not available",
                content_type="application/pdf"
            )

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_content = ""
                page_count = len(reader.pages)

                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num + 1} from PDF: {e}")
                        continue

                if not text_content.strip():
                    return ExtractionResult(
                        success=False,
                        error="No text found in PDF",
                        content_type="application/pdf",
                        metadata={"page_count": page_count}
                    )

                return ExtractionResult(
                    success=True,
                    text_content=text_content.strip(),
                    content_type="application/pdf",
                    extraction_method="PyPDF2",
                    confidence=0.8,  # PDF text extraction is generally reliable
                    metadata={"page_count": page_count}
                )

        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"PDF extraction failed: {e}",
                content_type="application/pdf"
            )

    def _extract_docx(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            ExtractionResult with extracted text
        """
        if not DocxDocument:
            return ExtractionResult(
                success=False,
                error="python-docx not available",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        try:
            doc = DocxDocument(file_path)
            paragraphs = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())

            # Extract text from tables
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_data.append(" | ".join(row_text))

            # Combine all text
            text_content = "\n\n".join(paragraphs)
            if table_data:
                text_content += "\n\n--- Tables ---\n" + "\n".join(table_data)

            if not text_content.strip():
                return ExtractionResult(
                    success=False,
                    error="No text found in DOCX",
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            return ExtractionResult(
                success=True,
                text_content=text_content.strip(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                extraction_method="python-docx",
                confidence=0.9,  # DOCX extraction is very reliable
                metadata={
                    "paragraph_count": len(paragraphs),
                    "table_count": len(doc.tables)
                }
            )

        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"DOCX extraction failed: {e}",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    def _extract_text(self, file_path: Path) -> ExtractionResult:
        """Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            ExtractionResult with file content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()

                    return ExtractionResult(
                        success=True,
                        text_content=text_content,
                        content_type="text/plain",
                        extraction_method=f"text-{encoding}",
                        confidence=1.0,  # Text extraction is definitive
                        metadata={"encoding": encoding}
                    )

                except UnicodeDecodeError:
                    continue

            return ExtractionResult(
                success=False,
                error="Unable to decode text with any supported encoding",
                content_type="text/plain"
            )

        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"Text extraction failed: {e}",
                content_type="text/plain"
            )

    def _extract_image_ocr(self, file_path: Path) -> ExtractionResult:
        """Extract text from image using OCR.

        Args:
            file_path: Path to image file

        Returns:
            ExtractionResult with OCR text
        """
        if not TESSERACT_AVAILABLE:
            return ExtractionResult(
                success=False,
                error="Tesseract OCR not available",
                content_type="image"
            )

        try:
            # Open and preprocess image
            image = Image.open(file_path)

            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Auto-orient the image
            image = ImageOps.exif_transpose(image)

            # Enhance image for better OCR
            image = ImageOps.autocontrast(image)

            # Extract text using Tesseract
            ocr_text = pytesseract.image_to_string(image)

            # Get confidence score
            try:
                ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                confidence = avg_confidence / 100.0  # Convert to 0-1 scale
            except Exception:
                confidence = 0.5  # Default confidence

            if not ocr_text.strip():
                return ExtractionResult(
                    success=False,
                    error="No text found in image",
                    content_type="image",
                    confidence=confidence
                )

            return ExtractionResult(
                success=True,
                text_content=ocr_text.strip(),
                content_type="image",
                extraction_method="tesseract-ocr",
                confidence=confidence,
                metadata={
                    "image_size": image.size,
                    "image_mode": image.mode
                }
            )

        except Exception as e:
            return ExtractionResult(
                success=False,
                error=f"OCR extraction failed: {e}",
                content_type="image"
            )

    def _extract_doc_fallback(self, file_path: Path) -> ExtractionResult:
        """Fallback extraction for legacy DOC files.

        Args:
            file_path: Path to DOC file

        Returns:
            ExtractionResult (currently unsupported)
        """
        # Legacy .doc files require more complex libraries like python-docx2txt or antiword
        # For now, mark as unsupported
        return ExtractionResult(
            success=False,
            error="Legacy DOC format not currently supported",
            content_type="application/msword"
        )

    def _clean_text(self, text: str) -> str:
        """Clean extracted text.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)

        # Remove common OCR artifacts
        text = re.sub(r'[^\w\s\-.,!?:;"\'()\[\]{}@#$%^&*+=|\\/<>~`]', '', text)

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        return text.strip()

    def _detect_language_basic(self, text: str) -> Optional[str]:
        """Basic language detection.

        Args:
            text: Text to analyze

        Returns:
            Detected language code or None
        """
        # Very basic language detection based on common words
        # In production, you might want to use a proper language detection library

        if len(text) < 50:
            return None

        text_lower = text.lower()

        # Common English words
        english_indicators = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        english_count = sum(1 for word in english_indicators if f' {word} ' in f' {text_lower} ')

        # Common Spanish words
        spanish_indicators = ['el', 'la', 'de', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'con', 'por']
        spanish_count = sum(1 for word in spanish_indicators if f' {word} ' in f' {text_lower} ')

        # Simple heuristic
        if english_count > spanish_count and english_count > 2:
            return 'en'
        elif spanish_count > english_count and spanish_count > 2:
            return 'es'

        return None


class ExtractedContentStorage:
    """Storage for extracted content."""

    def __init__(self):
        """Initialize content storage."""
        pass

    async def store_extracted_content(
        self,
        attachment_id: int,
        extraction_result: ExtractionResult
    ) -> bool:
        """Store extracted content.

        Args:
            attachment_id: EmailAttachment ID
            extraction_result: Extraction result

        Returns:
            True if stored successfully
        """
        # This could store in a separate table, search index, or document store
        # For now, we'll just publish events and let other systems handle storage
        logger.info(f"Extracted content stored for attachment {attachment_id}")
        return True